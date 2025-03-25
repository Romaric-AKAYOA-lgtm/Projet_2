from django.shortcuts import render, get_object_or_404, redirect
from django.utils.timezone import now
from Activation.models import Activation
from secretaire.views import get_username_from_session
from .models import ProgrammeVisite
from .forms import ProgrammeVisiteForm
from visite.models import Visite
from secretaire.models import Secretaire
from datetime import date
from reportlab.lib.pagesizes import letter, landscape, A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Spacer, Paragraph
from django.http import HttpResponse
from reportlab.lib import colors
import io
from io import BytesIO
from django.db.models import Q


def programme_visite_list(request):
    """Affiche la page d'accueil avec la gestion du personnel et la vérification d'activation."""
    
    # 🔹 Vérifier l'activation
    activation = Activation.objects.first()
    if not activation or not activation.is_valid():
        return redirect("Activation:activation_page")  # Redirige vers une page d'activation si expiré
    username = get_username_from_session(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    # Filtrer les programmes de visite pour afficher seulement ceux qui ont un statut confirmé
    programmes = ProgrammeVisite.objects.filter(visite__statut='confirmé').order_by('-date_creation')
    return render(request, 'programme_visite/list.html', {'programmes': programmes , 'username':username # Passer le nom d'utilisateur dans le contexte
 })


def programme_visite_create(request):
    username = get_username_from_session(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    # Obtenir la date système (aujourd'hui)
    today = now().date()
    
    # Filtrer les visites confirmées ayant une date de visite <= aujourd'hui 
    # et qui ne sont pas encore enregistrées dans ProgrammeVisite
    visites = Visite.objects.filter(
        statut='confirmé',
        date_visite__lte=today  # Sélectionner les visites dont la date est <= aujourd'hui
    ).exclude(id__in=ProgrammeVisite.objects.values_list('visite_id', flat=True)).order_by('-date_visite')

    # Récupérer les secrétaires avec une date de début non nulle, triées par date de début décroissante
    secretaires = Secretaire.objects.filter(date_debut__isnull=False).order_by('-date_debut')

    form = ProgrammeVisiteForm(request.POST or None)

    if request.method == "POST" and form.is_valid():
        form.save()
        return redirect("programme_visite:ProgrammeVisite")

    return render(request, "programme_visite/form.html", {
        'username':username,
        "form": form,
        "visites": visites,
        "secretaires": secretaires,  # Passez les secrétaires au contexte
    })
    
def programme_visite_detail(request, id):
    username = get_username_from_session(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    programme = get_object_or_404(ProgrammeVisite, id=id)
    return render(request, 'programme_visite/detail.html', {'programme': programme, 'username':username})

def modifier_programme_visite(request, id):
    username = get_username_from_session(request)
    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session
    # Récupère le programme de visite par son ID
    programme = get_object_or_404(ProgrammeVisite, id=id)
    # Récupère toutes les visites confirmées
    visites = Visite.objects.filter(statut='confirmé').order_by('-date_visite')
    # Récupère toutes les secrétaires ayant une date de début
    secretaires = Secretaire.objects.filter(date_debut__isnull=False).order_by('-date_debut')
    if request.method == "POST":
        form = ProgrammeVisiteForm(request.POST, instance=programme)
        
        if form.is_valid():
            # Sauvegarder les modifications si le formulaire est valide
            form.save()
            return redirect('programme_visite:ProgrammeVisite')
        else:
            # Si le formulaire n'est pas valide, il contiendra des erreurs
            pass  # Les erreurs seront gérées dans le template
    else:
        form = ProgrammeVisiteForm(instance=programme)

    # Rendre le template avec les données
    return render(request, 'programme_visite/ajouter_programme_visite.html', {
        'username': username,
        'form': form,  # Transmettre le formulaire avec ses erreurs si elles existent
        'programme_visite': programme,
        'visites': visites,
        'secretaires': secretaires,  # Ajout de la secrétaire récente
    })


def supprimer_programme_visite(request, id):
    programme = get_object_or_404(ProgrammeVisite, id=id)
    programme.delete()
    return redirect('programme_visite:ProgrammeVisite')

def programme_visite_search(request):
    username = get_username_from_session(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    # Récupérer les paramètres de recherche
    query = request.GET.get('query', '')  # Recherche globale
    critere = request.GET.get('criteres', '')  # Critère de recherche sélectionné (visite, secretaire, statut, etc.)
    
    # Initialisation de la queryset avec tous les programmes de visite
    programmes = ProgrammeVisite.objects.all()

    # Si un critère et un terme de recherche sont saisis, filtrer en fonction du critère
    if critere and query:
        if critere == 'visite':
            programmes = programmes.filter(visite__visiteur__nom__icontains=query)  # Filtrer par nom du visiteur
        elif critere == 'secretaire':
            programmes = programmes.filter(secretaire__last_name__icontains=query)  # Assurez-vous que 'last_name' est un champ de Secretaire
        elif critere == 'statut':
            programmes = programmes.filter(statut__icontains=query)
        elif critere == 'date_creation':
            programmes = programmes.filter(date_creation__icontains=query)  # Recherche par date
    elif query:
        # Si un terme de recherche est saisi sans critère, effectuer une recherche globale sur plusieurs champs
        programmes = programmes.filter(
            Q(visite__visiteur__nom__icontains=query) |  # Recherche sur le nom du visiteur
            Q(secretaire__last_name__icontains=query) |
            Q(statut__icontains=query) |
            Q(date_creation__icontains=query)
        )

    # Ajouter des filtres supplémentaires pour les dates (date_debut et date_fin)
    date_debut = request.GET.get('date_debut', '')
    date_fin = request.GET.get('date_fin', '')
    
    # Appliquer les filtres supplémentaires si les champs sont remplis
    if date_debut:
        programmes = programmes.filter(date_creation__gte=date_debut)  # Filtrer après la date_debut
    if date_fin:
        programmes = programmes.filter(date_creation__lte=date_fin)  # Filtrer avant la date_fin

    # Si aucun résultat n'est trouvé, on peut afficher un message ou laisser la queryset vide
    if not programmes:
        programmes = None  # Pas de programmes trouvés

    # Retourner la réponse au template avec les données filtrées
    return render(request, 'programme_visite/search.html', {
        'username':username,
        'programmes': programmes,
        'query': query,
        'criteres': critere,
        'date_debut': date_debut,
        'date_fin': date_fin,
    })

import io
from datetime import date, timedelta
from django.http import HttpResponse
from django.shortcuts import redirect
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from .models import ProgrammeVisite
from secretaire.models import Secretaire
from Activation.models import Activation

def obtenir_secretaire_recente():
    """Récupère la secrétaire la plus récente."""
    return Secretaire.objects.order_by('-date_debut').first()

def get_username_from_session(request):
    """Récupère le nom d'utilisateur à partir de la session."""
    return request.session.get('username')

from datetime import date, timedelta

def imprimer_tous_les_programmes(request):
    """Génère un PDF des programmes du jour, de la semaine et du mois."""
    date_aujourdhui = date.today()
    
    # Calcul de la date de début de semaine (sans l'heure)
    date_debut_semaine = date_aujourdhui - timedelta(days=6)
    
    # Calcul de la date de début de mois (sans l'heure)
    date_debut_mois = date_aujourdhui.replace(day=1)
    
    # Filtrage des programmes selon la date sans tenir compte de l'heure
    programmes_jour = ProgrammeVisite.objects.filter(date_creation__date=date_aujourdhui)
    programmes_semaine = ProgrammeVisite.objects.filter(date_creation__date__range=[date_debut_semaine, date_aujourdhui])
    programmes_mois = ProgrammeVisite.objects.filter(date_creation__date__range=[date_debut_mois, date_aujourdhui])

    secretaire_recente = obtenir_secretaire_recente()
    return generer_word(request, programmes_jour, programmes_semaine, programmes_mois, "PROGRAMMES DE VISITE", secretaire_recente)


from io import BytesIO
from datetime import date
from django.http import HttpResponse
from docx import Document
from io import BytesIO
from datetime import date
from django.http import HttpResponse
from docx import Document
from docx.shared import Inches
from io import BytesIO
from datetime import date
from django.http import HttpResponse
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import nsmap
from io import BytesIO
from datetime import date
from django.http import HttpResponse
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import nsmap

def generer_word(request, programmes_jour, programmes_semaine, programmes_mois, titre, secretaire_recente):
    """Génère un fichier Word des programmes en orientation paysage sans en-tête et avec cadrage des lignes du tableau."""
    activation = Activation.objects.first()
    if not activation or not activation.is_valid():
        return redirect("Activation:activation_page")
    
    username = get_username_from_session(request)
    if not username:
        return redirect('login')
    
    doc = Document()
    
    # Définir l'orientation paysage
    section = doc.sections[0]
    section.orientation = 1  # Paysage
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    
    # Centrer le titre du document
    titre_paragraph = doc.add_paragraph(titre, style='Heading1')
    titre_paragraph.alignment = 1  # Centrer
    
    def ajouter_section(nom_section, programmes):
        doc.add_paragraph("\n")
        section_paragraph = doc.add_paragraph(nom_section, style='Heading2')
        section_paragraph.alignment = 1  # Centrer
        
        if not programmes:
            doc.add_paragraph("Aucun programme disponible.", style='Normal')
            return
        
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'  # Ajout du cadrage des lignes
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "ID"
        hdr_cells[1].text = "Visite"
        hdr_cells[2].text = "Secrétaire"
        hdr_cells[3].text = "Statut"
        hdr_cells[4].text = "Date création"
        hdr_cells[5].text = "Heure Début"
        hdr_cells[6].text = "Heure Fin"
        
        for programme in programmes:
            secretaire_complete = f"{programme.secretaire.last_name} {programme.secretaire.first_name}" if programme.secretaire else "N/A"
            
            # Conversion de date_creation en chaîne si c'est un datetime
            date_creation_str = programme.date_creation.strftime("%d/%m/%Y") if programme.date_creation else "N/A"
            
            # Création de la ligne du tableau
            row_cells = table.add_row().cells
            row_cells[0].text = str(programme.id)
            row_cells[1].text = str(programme.visite)
            row_cells[2].text = secretaire_complete
            row_cells[3].text = programme.statut
            row_cells[4].text = date_creation_str  # Ajout de la date formatée
            row_cells[5].text = programme.heure_debut.strftime("%H:%M") if programme.heure_debut else "N/A"
            row_cells[6].text = programme.heure_fin.strftime("%H:%M") if programme.heure_fin else "N/A"

    ajouter_section("Programmes du Jour", programmes_jour)
    ajouter_section("Programmes de la Semaine", programmes_semaine)
    ajouter_section("Programmes du Mois", programmes_mois)
    
    # Ajouter la signature alignée à droite
    doc.add_paragraph("\n")
    if secretaire_recente:
        lieu_date = f"Brazzaville, Fait le : {date.today().strftime('%d/%m/%Y')}"
        signataire = f"{getattr(secretaire_recente, 'fonction', 'Secrétaire')}, {secretaire_recente.last_name} {secretaire_recente.first_name}"
        
        p1 = doc.add_paragraph()
        p1.alignment = 2  # Alignement à droite
        p1.add_run(lieu_date)
        
        doc.add_paragraph("\n\n\n")
        
        p2 = doc.add_paragraph()
        p2.alignment = 2  # Alignement à droite
        p2.add_run(signataire)
    
    # Sauvegarde du fichier en mémoire
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="programmes_visite.docx"'
    return response
